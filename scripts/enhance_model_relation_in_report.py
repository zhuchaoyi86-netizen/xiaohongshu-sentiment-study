from pathlib import Path

import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/xinxinhuashe/Documents/1/python课")
BASE = ROOT / "小组共享材料" / "基于文本挖掘的小红书食品种草笔记情感对点赞热度的影响研究_补充模型关系版.docx"
ORIGINAL = Path("/Users/xinxinhuashe/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_ah3md0gkwwsx22_bc68/temp/drag/基于文本挖掘的小红书食品种草笔记情感对点赞热度的影响研究.docx")
OUT = ROOT / "小组共享材料" / "基于文本挖掘的小红书食品种草笔记情感对点赞热度的影响研究_模型关系增强版.docx"
FIG = ROOT / "第四部分材料" / "图" / "报告_变量之间的模型关系图.png"


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def style_run(run, size=10.5, bold=False, color=None):
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def make_relation_figure():
    matplotlib.rcParams["font.sans-serif"] = ["Arial Unicode MS", "PingFang SC", "Heiti SC", "SimHei"]
    matplotlib.rcParams["axes.unicode_minus"] = False

    FIG.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8.4, 4.6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")

    def box(x, y, w, h, text, fc, ec="#8a8a8a", color="#1f3347", fs=12):
        rect = plt.Rectangle((x, y), w, h, linewidth=1.2, edgecolor=ec, facecolor=fc, zorder=2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, color=color, weight="bold")

    box(4.0, 2.45, 2.0, 1.05, "点赞热度\nlog_likes", "#ffe6dc", "#e8534a", "#c03d35", 13)
    box(0.75, 4.25, 2.15, 0.8, "情绪密度\nsentiment_density", "#def4f1", "#2d908b", "#166b68", 11)
    box(0.75, 2.7, 2.15, 0.8, "情感得分\nsentiment_score", "#e2ecf9", "#4b7cad", "#17314d", 11)
    box(0.75, 1.15, 2.15, 0.8, "混合评价\nmixed_review", "#ffe0dc", "#e8534a", "#c03d35", 11)
    box(4.0, 0.55, 2.0, 0.75, "控制变量", "#f5f1ea", "#9b9287", "#374151", 11)
    box(7.3, 3.75, 2.0, 0.75, "图文呈现\n配图数量", "#e8f2e1", "#7a9c72", "#3d6b45", 10)
    box(7.3, 2.35, 2.0, 0.75, "文本形式\n长度/标签/标点", "#fff0da", "#ee8b42", "#ad5d21", 10)
    box(7.3, 0.95, 2.0, 0.75, "发布时间\n上午/下午/晚上", "#e2ecf9", "#4b7cad", "#17314d", 10)

    def arrow(x1, y1, x2, y2, color, label):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", lw=2, color=color))
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.15, label, color=color, fontsize=11, weight="bold")

    arrow(2.9, 4.65, 4.0, 3.25, "#2d908b", "+")
    arrow(2.9, 3.1, 4.0, 3.0, "#4b7cad", "±")
    arrow(2.9, 1.55, 4.0, 2.65, "#e8534a", "-")
    arrow(6.0, 2.95, 7.3, 4.1, "#7a9c72", "+")
    arrow(6.0, 2.95, 7.3, 2.72, "#ee8b42", "控制")
    arrow(6.0, 2.95, 7.3, 1.32, "#4b7cad", "控制")

    ax.text(
        0.75,
        5.55,
        "图  模型关系框架：情感表达、评价结构与平台情境共同影响点赞热度",
        fontsize=13,
        weight="bold",
        color="#17314d",
    )
    ax.text(
        0.75,
        0.18,
        "说明：绿色/蓝色箭头表示正向或辅助解释，红色箭头表示负向关系；控制变量用于降低图文、文本形式和发布时间差异带来的干扰。",
        fontsize=9.5,
        color="#6b7280",
    )
    plt.tight_layout()
    fig.savefig(FIG, dpi=220, bbox_inches="tight")
    plt.close(fig)


def add_paragraph_before(target, text, style="Normal"):
    para = target.insert_paragraph_before(text)
    para.style = style
    return para


def main():
    make_relation_figure()
    src = BASE if BASE.exists() else ORIGINAL
    doc = Document(str(src))

    relation_idx = None
    robust_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "变量之间的模型关系" and relation_idx is None:
            relation_idx = i
        if t == "鲁棒性检验" and relation_idx is not None:
            robust_idx = i
            break

    if relation_idx is None or robust_idx is None:
        raise ValueError("未找到“变量之间的模型关系”或“鲁棒性检验”标题。")

    # 删除原变量关系标题和鲁棒性标题之间的旧段落，保留标题本身。
    for p in list(doc.paragraphs[relation_idx + 1 : robust_idx]):
        delete_paragraph(p)

    robust_para = None
    for p in doc.paragraphs:
        if p.text.strip() == "鲁棒性检验":
            robust_para = p
            break

    new_paragraphs = [
        "在本文的实证模型中，变量之间并不是孤立进入回归方程的，而是共同构成了“情感表达—评价结构—平台情境—点赞热度”的作用链条。被解释变量为缩尾并取对数后的点赞热度 log_likes_winsorized，核心解释变量包括情绪密度 sentiment_density、缩尾后的情感得分 sentiment_score_winsorized 以及混合评价变量 mixed_review_flag；同时，模型控制文本长度、配图数量、话题标签数、标点语气、营销话术和发布时间段等变量，以减少非情感因素对估计结果的干扰。",
        "首先，情绪密度用于衡量单位文本长度中的情绪表达集中程度。与单纯统计正向词数量相比，情绪密度更能反映一篇笔记在有限篇幅内是否形成了鲜明态度。基准回归结果显示，情绪密度系数为正，说明在控制图文形式和发布时间后，情绪表达越集中、越鲜明的食品种草笔记整体上越容易获得点赞。其次，情感得分在与情绪密度同时进入模型后并未表现出稳定的正向作用，这提示本文不能简单得出“正向词越多点赞越高”的结论，而应更重视情绪表达方式本身。",
        "再次，混合评价变量刻画的是文本中是否同时包含推荐与保留意见、种草与避雷等正负并存的信息结构。模型结果表明，混合评价对点赞热度具有稳定的负向影响。这说明当食品种草笔记的评价态度不够单一、同时传递积极和消极信息时，用户可能会形成更加谨慎的判断，从而降低点赞意愿。最后，配图数量和发布时间变量说明，小红书食品笔记的点赞热度还受到平台内容形态和用户活跃时段影响，因此情感变量的作用需要放在图文呈现和发布时间环境中共同理解。",
    ]
    for para_text in new_paragraphs:
        add_paragraph_before(robust_para, para_text)

    cap = add_paragraph_before(robust_para, "图 3  变量之间的模型关系框架")
    cap.style = doc.styles["Caption"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        style_run(run, size=10, bold=True)

    fig_para = robust_para.insert_paragraph_before()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run().add_picture(str(FIG), width=Inches(5.9))

    note = add_paragraph_before(
        robust_para,
        "从图 3 可以看出，情绪密度、情感得分和混合评价共同构成文本情感层面的解释变量，其中情绪密度主要反映情绪表达的集中程度，混合评价则反映评价结构的复杂性；图文呈现、文本形式和发布时间作为控制变量进入模型，用于帮助识别情感变量对点赞热度的相对独立影响。",
    )
    note.style = doc.styles["Normal"]

    table_title = add_paragraph_before(robust_para, "表  模型变量关系说明")
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in table_title.runs:
        style_run(run, size=10, bold=True)

    table_data = [
        ["变量层级", "代表变量", "与点赞热度的关系", "研究含义"],
        ["因变量", "log_likes_winsorized", "被解释对象", "衡量笔记点赞热度，并降低极端爆款影响"],
        ["核心解释变量", "sentiment_density", "正向关系", "情绪表达越集中，越容易引发点赞互动"],
        ["核心解释变量", "sentiment_score_winsorized", "需结合密度解释", "单纯情感得分并不等同于更高点赞"],
        ["情感结构变量", "mixed_review_flag", "负向关系", "正负评价并存会削弱用户点赞意愿"],
        ["控制变量", "imagenumber / text_length / time_period", "辅助解释", "控制图文呈现、文本信息量和发布时间差异"],
    ]

    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = "Table Grid"
    for i, row in enumerate(table_data):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 3 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    style_run(run, size=9.5, bold=(i == 0), color=RGBColor(255, 255, 255) if i == 0 else None)
            if i == 0:
                cell._tc.get_or_add_tcPr()
                shading = cell._tc.get_or_add_tcPr()
                # Avoid direct XML shading dependency where possible; table remains readable without fill.
    robust_para._p.addprevious(table._tbl)

    # 新增模型关系图后，后续鲁棒性图的编号顺延，避免报告中出现重复“图 3”。
    replacements = {
        "图 3  情绪密度系数在不同设定下的稳健性": "图 4  情绪密度系数在不同设定下的稳健性",
        "图 3所示": "图 4所示",
        "图 3也揭示": "图 4也揭示",
        "图 4展示": "图 5展示",
        "图 4  逐步加入控制变量时情绪密度系数的变化": "图 5  逐步加入控制变量时情绪密度系数的变化",
        "图 5所示": "图 6所示",
        "图 5  混合评价系数在不同设定下的稳健性": "图 6  混合评价系数在不同设定下的稳健性",
    }
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)
    print(FIG)


if __name__ == "__main__":
    main()
