from __future__ import annotations

from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path("/Users/xinxinhuashe/Documents/1/python课")
OUTPUT_DIR = BASE_DIR / "output/docx"
DOCX_PATH = OUTPUT_DIR / "第三部分_数据预处理与变量构建.docx"

PART3_MD = BASE_DIR / "notes/第三部分正式成稿.md"
SCRIPT_PATHS = [
    BASE_DIR / "scripts/clean_xiaohongshu_xlsx.py",
    BASE_DIR / "scripts/filter_food_seed_notes.py",
    BASE_DIR / "scripts/add_sentiment_variables.py",
]
SUMMARY_TABLES = [
    ("表 3-1 数据清洗汇总", BASE_DIR / "output/cleaned_xiaohongshu/小红书商品笔记_清洗汇总.csv"),
    ("表 3-2 食品种草笔记筛选汇总", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_筛选汇总.csv"),
    ("表 3-3 食品种草笔记情感分布汇总", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_情感分布汇总.csv"),
    ("表 3-4 情感方向变量与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_情感方向.csv"),
    ("表 3-5 情绪强度与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_情绪强度.csv"),
    ("表 3-6 情绪密度与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_情绪密度.csv"),
]

APPENDIX_TABLES = [
    ("附表 A-1 混合评价与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_混合评价.csv"),
    ("附表 A-2 情感结构变量与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_情感结构.csv"),
    ("附表 A-3 配图数量与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_配图数量.csv"),
    ("附表 A-4 发布时间时段与点赞热度初步分析", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_发布时间.csv"),
    ("附表 A-5 点赞热度变量说明", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_点赞热度变量说明.csv"),
    ("附表 A-6 变量与初步结论对应表", BASE_DIR / "output/food_seed_notes/小红书食品种草笔记_变量分析_变量结论对应表.csv"),
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_page_break_before(paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    page_break_before = OxmlElement("w:pageBreakBefore")
    ppr.append(page_break_before)


def apply_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def apply_heading_style(style, size: int, east_asia_font: str = "黑体") -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    pf = style.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    set_east_asia_font(run, "黑体")


def add_meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(11)
    set_east_asia_font(run, "宋体")


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_east_asia_font(run, "宋体")


def read_csv_rows(path: Path) -> list[list[str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.reader(f))


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    set_east_asia_font(run, "宋体")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体")


def add_table_from_csv(doc: Document, caption: str, path: Path) -> None:
    rows = read_csv_rows(path)
    if not rows:
        return
    add_caption(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_soft_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break()


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "等宽更纱黑体 SC")


def parse_markdown_sections(md_text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_heading = ""
    current_lines: list[str] = []
    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if line.startswith("## "):
            if current_heading:
                sections.append((current_heading, "\n".join(current_lines).strip()))
            current_heading = line[3:].strip()
            current_lines = []
        elif line.startswith("# "):
            continue
        else:
            current_lines.append(raw_line)
    if current_heading:
        sections.append((current_heading, "\n".join(current_lines).strip()))
    return sections


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    apply_normal_style(doc)
    apply_heading_style(doc.styles["Heading 1"], 14)
    apply_heading_style(doc.styles["Heading 2"], 12)

    add_title(doc, "第三部分 数据预处理与变量构建")
    add_meta(doc, "题目：基于文本挖掘的小红书食品种草笔记情感对点赞热度的影响研究")

    md_text = PART3_MD.read_text(encoding="utf-8")
    for heading, body in parse_markdown_sections(md_text):
        p = doc.add_paragraph(style="Heading 1")
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(heading)
        set_east_asia_font(run, "黑体")
        for para in [x.strip() for x in body.split("\n\n") if x.strip()]:
            add_body_paragraph(doc, para.replace("`", ""))
        if heading == "3.1 样本筛选与原始数据清洗":
            add_table_from_csv(doc, SUMMARY_TABLES[0][0], SUMMARY_TABLES[0][1])
            add_table_from_csv(doc, SUMMARY_TABLES[1][0], SUMMARY_TABLES[1][1])
        if heading == "3.2 文本清洗与情感量化处理":
            add_table_from_csv(doc, SUMMARY_TABLES[2][0], SUMMARY_TABLES[2][1])
        if heading == "3.5 基于构建变量的初步数据分析":
            add_table_from_csv(doc, SUMMARY_TABLES[3][0], SUMMARY_TABLES[3][1])
            add_table_from_csv(doc, SUMMARY_TABLES[4][0], SUMMARY_TABLES[4][1])
            add_table_from_csv(doc, SUMMARY_TABLES[5][0], SUMMARY_TABLES[5][1])

    appendix_heading = doc.add_paragraph(style="Heading 1")
    set_paragraph_page_break_before(appendix_heading)
    appendix_heading.paragraph_format.first_line_indent = Cm(0)
    run = appendix_heading.add_run("附录 A：补充表与脚本")
    set_east_asia_font(run, "黑体")

    appendix_intro = "为避免正文中表格过于集中，部分补充变量分析表与脚本统一置于附录。"
    add_body_paragraph(doc, appendix_intro)

    appendix_table_heading = doc.add_paragraph(style="Heading 2")
    appendix_table_heading.paragraph_format.first_line_indent = Cm(0)
    appendix_table_run = appendix_table_heading.add_run("A.1 补充变量分析表")
    set_east_asia_font(appendix_table_run, "黑体")
    for caption, path in APPENDIX_TABLES:
        add_table_from_csv(doc, caption, path)

    add_soft_page_break(doc)
    script_heading = doc.add_paragraph(style="Heading 2")
    script_heading.paragraph_format.first_line_indent = Cm(0)
    script_run = script_heading.add_run("A.2 数据清洗与变量构建脚本")
    set_east_asia_font(script_run, "黑体")
    for script_path in SCRIPT_PATHS:
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(script_path.name)
        set_east_asia_font(run, "黑体")
        code_text = script_path.read_text(encoding="utf-8")
        for line in code_text.splitlines():
            add_code_paragraph(doc, line)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
